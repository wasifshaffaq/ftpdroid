import sys
import subprocess

def install_package(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package, "--quiet"])

try:
    import docx
except ImportError:
    install_package('python-docx')
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('FTPDroid - Session Changes Summary', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This document outlines in extreme detail every change made to the FTPDroid application during this session to fix UI/Backend desynchronization, progress tracking issues, and queue management bugs.')

doc.add_heading('1. Data Synchronization & Event Emission (Backend)', level=1)
p1 = doc.add_paragraph()
p1.add_run('File: ').bold = True
p1.add_run('app/src/main/kotlin/com/ftpdroid/app/data/network/ftp/FtpServerManager.kt\n')
p1.add_run('• Issue: ').bold = True
p1.add_run('The background FTP server (LoggingFtplet) was dropping transfer status events when the system was busy because it used a non-blocking `tryEmit` call on a limited buffer SharedFlow. This caused the UI to permanently miss "TransferFinished" signals, resulting in stuck transfers.\n')
p1.add_run('• Fix: ').bold = True
p1.add_run('Upgraded the event emission from `tryEmit` to a guaranteed `emit` inside a CoroutineScope. This ensures backpressure is handled and no events are lost.\n')
p1.add_run('• Thread Safety: ').bold = True
p1.add_run('Replaced standard `MutableMap` with `ConcurrentHashMap` for `startTimes` and `activeTransfers` to prevent concurrency crashes when multiple clients connect or transfer simultaneously.\n')
p1.add_run('• Physical Path Resolution: ').bold = True
p1.add_run('Modified the `STOR`, `APPE`, and `RETR` FTP commands to correctly resolve the physical absolute path of files on the device instead of relying on the virtual FTP path. This is crucial for the file size polling mechanism to find the file on disk.\n')
p1.add_run('• Accurate Final Size: ').bold = True
p1.add_run('Improved the `TransferFinished` collector. If the FTP session reports 0 bytes for an upload upon completion, the manager now falls back to checking the actual physical file size on the disk before saving the final record to the database.\n')

doc.add_heading('2. Sequential Queue Processing & Speed Calculation (Backend)', level=1)
p2 = doc.add_paragraph()
p2.add_run('File: ').bold = True
p2.add_run('app/src/main/kotlin/com/ftpdroid/app/service/TransferService.kt\n')
p2.add_run('• Issue: ').bold = True
p2.add_run('The app was attempting to execute all transfers in parallel. This caused network bottlenecks, corrupted the single `FtpClientManager` instance, and caused speed calculations to freeze at 0 B/s.\n')
p2.add_run('• Sequential Queue Fix: ').bold = True
p2.add_run('Completely refactored the service to use a sequential processing loop. It now uses a Coroutine `Mutex` and processes transfers one-by-one. Pending transfers sit in a "QUEUED" state until the active "IN_PROGRESS" transfer completes or fails.\n')
p2.add_run('• Speed Calculation: ').bold = True
p2.add_run('Replaced the hardcoded or inaccurate speed metrics with a delta-based calculation. The service now checks the bytes transferred every 1000ms (1 second), compares it to the previous value, and calculates the exact bytes-per-second speed, persisting this to the database.\n')

doc.add_heading('3. Queue Ordering & Database Logic (Database)', level=1)
p3 = doc.add_paragraph()
p3.add_run('File: ').bold = True
p3.add_run('app/src/main/kotlin/com/ftpdroid/app/data/local/db/dao/TransferDao.kt\n')
p3.add_run('• Issue: ').bold = True
p3.add_run('Files in the transfer queue would randomly swap positions due to unstable database sorting logic.\n')
p3.add_run('• Fix: ').bold = True
p3.add_run('Updated the `getPendingTransfers` SQL query to use deterministic sorting: `ORDER BY startedAt ASC, id ASC`. This ensures the queue is strictly FIFO (First-In-First-Out) and visually stable.\n')

doc.add_heading('4. UI State Recycling & Visual Glitches (Frontend)', level=1)
p4 = doc.add_paragraph()
p4.add_run('File: ').bold = True
p4.add_run('app/src/main/kotlin/com/ftpdroid/app/ui/screen/transfer/TransferQueueScreen.kt\n')
p4.add_run('• Issue: ').bold = True
p4.add_run('The "100% progress thumb with 0% fill" bug was caused by Jetpack Compose `LazyColumn` recycling state from old, completed list items into new, pending items.\n')
p4.add_run('• Fix: ').bold = True
p4.add_run('Added a unique identity key (`key = { it.id }`) to the `items` builder in the `LazyColumn`. This forces Compose to treat each transfer as a distinct element, preventing ghost animations and state bleeding.\n')
p4.add_run('• ETA Display: ').bold = True
p4.add_run('Modified the ETA display logic. If the transfer speed is 0 B/s (e.g., during initialization or a stall), the UI now explicitly shows "Calculating..." instead of a blank space or crashing due to a divide-by-zero error.\n')

doc.add_heading('5. SFTP Progress Tracking (Backend)', level=1)
p5 = doc.add_paragraph()
p5.add_run('File: ').bold = True
p5.add_run('app/src/main/kotlin/com/ftpdroid/app/data/network/ftp/SftpClientManager.kt\n')
p5.add_run('• Issue: ').bold = True
p5.add_run('SFTP transfers had no progress tracking implemented.\n')
p5.add_run('• Fix: ').bold = True
p5.add_run('Implemented `net.schmizz.sshj.xfer.TransferListener` and `StreamCopier.Listener` during `uploadFile` and `downloadFile` operations to periodically report the transferred bytes to the UI.\n')

doc.add_heading('6. FTP Client Stability (Backend)', level=1)
p6 = doc.add_paragraph()
p6.add_run('File: ').bold = True
p6.add_run('app/src/main/kotlin/com/ftpdroid/app/data/network/ftp/FtpClientManager.kt\n')
p6.add_run('• Issue: ').bold = True
p6.add_run('Fetching file sizes after opening the data stream could cause the control connection to hang.\n')
p6.add_run('• Fix: ').bold = True
p6.add_run('Moved the `ftpClient.listFiles(remotePath)` call to retrieve the file size *before* the `retrieveFileStream` call is made in `downloadFile`. \n')

doc.add_heading('7. Missing Value Placeholders (Frontend)', level=1)
p7 = doc.add_paragraph()
p7.add_run('Files: ').bold = True
p7.add_run('TransferQueueViewModel.kt, TransferHistoryViewModel.kt\n')
p7.add_run('• Issue: ').bold = True
p7.add_run('Files sent to the phone via the local FTP Server showed "Profile -1" as the source.\n')
p7.add_run('• Fix: ').bold = True
p7.add_run('Added formatting logic to display "Incoming (FTP Server)" whenever the `profileId` is -1, making the origin of the file clear to the user.\n')

doc.save('FTPDroid_Changes_Summary.docx')
print("Successfully generated FTPDroid_Changes_Summary.docx")
